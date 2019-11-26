#!/usr/bin/env python3

# Importação de bibliotecas
import pandas as pd
import io
import os
import csv
import copy
import textract
from docx import Document

# Cria a tabela padrão
def gera_tabela():
    tabela = [("Café da Manhã","Chá"),
              ("Café da Manhã","Pão"),
              ("Café da Manhã","Fruta"),
              ("Almoço","Prato Principal"),
              ("Almoço","Guarnição"),
              ("Almoço","Sobremesa"),
              ("Almoço","Saladas"),
              ("Almoço","Opção Vegana"),
              ("Jantar","Prato Principal"),
              ("Jantar","Guarnição"),
              ("Jantar","Saladas"),
              ("Jantar","Opção Vegana"),
             ]
    tabela = pd.MultiIndex.from_tuples(tabela, names=["Refeição","NOME_MELHOR"])
    tabela = pd.DataFrame(columns=tabela)
    return tabela;

# Encontra o ano e o mês do arquivo
def gera_data(df):
    meses = ["erro",
             "janeiro",
             "fevereiro",
             "maio",
             "abril",
             "março",
             "junho",
             "julho",
             "agosto",
             "setembro",
             "outubro",
             "novembro",
             "dezembro",
            ]
    data = df.columns.values[0].split(" ")
    data = data[2] + "-"
    encontrou = False
    i = 1
    while (i <= len(meses)) & (not encontrou):
        if df.columns.values[0].split(" ")[1].lower() == meses[i]:
            if i < 10:
                data = data + "0" + str(i)
                encontrou = True
            else:
                data = data + str(i)
                encontrou = True
        i += 1
    if not encontrou:
        data = data + str(meses[0])
    return data;

# Função auxiliar da ordenação
def ordena(df,Refeição,lista):
    if lista:
        max = len(df)
    else:
        max = 1
    for i in range(0,max):
        if lista:
            string = df[i].split(" ")
        else:
            string = df.split(" ")
        for string_i in range(0,len(string)):
            if   string[string_i].lower() == "chá":
                return 0;
            elif string[string_i].lower() == "pão":
                return 1;
            elif string[string_i].lower() == "bife":
                if Refeição == "Almoço":
                    return 3;
                else:
                    return 8;
            elif string[string_i].lower() == "gelatina":
                    return 5;
            elif string[string_i].lower() == "alface":
                if Refeição == "Almoço":
                    return 6;
                else:
                    return 10;
            elif string[string_i].lower() == "pts":
                if Refeição == "Almoço":
                    return 7;
                else:
                    return 11;
    if   Refeição == "Café da Manhã":
        return 2;
    elif Refeição == "Almoço":
        return 4;
    else:
        return 9;

# Verifica se o input pode ser representado como um inteiro
def RepresentsInt(s):
    try: 
        int(s)
        return True
    except ValueError:
        return False

# Verifica se o input pode receber a função split()
def RepresentsStr(s):
    try: 
        s.split()
        return True
    except:
        return False

# Transforma docx em dataframe
def read_docx_tables(filename, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

# Function to insert row in the dataframe 
def Insert_row(row_number, df, row_value): 
    # Starting value of upper half 
    start_upper = 0
   
    # End value of upper half 
    end_upper = row_number 
   
    # Start value of lower half 
    start_lower = row_number 
   
    # End value of lower half 
    end_lower = df.shape[0] 
   
    # Create a list of upper_half index 
    upper_half = [*range(start_upper, end_upper, 1)] 
   
    # Create a list of lower_half index 
    lower_half = [*range(start_lower, end_lower, 1)] 
   
    # Increment the value of lower half by 1 
    lower_half = [x.__add__(1) for x in lower_half] 
   
    # Combine the two lists 
    index_ = upper_half + lower_half 
   
    # Update the index of the dataframe 
    df.index = index_ 
   
    # Insert a row at the end 
    df.loc[row_number] = row_value 
    
    # Sort the index labels 
    df = df.sort_index() 
   
    # return the dataframe 
    return df

# Corrige os casos particulares
def casos_particulares(df,caso):
    if caso == "05 Cardapio maio 2019 RU - FINAL com símbolos e descrições preparações.docx":
        df.iloc[1,2] = "01\nChá Maçã\nPão hot dogc/ doce de banana\n Laranja"
        df.iloc[1,3] = "02\nChá de morango\nPão francês presunto\nBanana"
        df.iloc[1,4] = "03\nChá de erva doce\nPão hot dog com pate de azeitona\nSalada de frutas"
        df.iloc[1,5] = "04\nChá de hortelã\nPão francês com queijo\nMamão"
        df.iloc[1,6] = "05\nChá de erva cidreira\nPão francês com presunto\nBanana"
    elif caso == "CARDÁPIO MARÇO 2019 - simbolos.docx":
        for i in range(0,4):
            df.iloc[0,i] = "n\no\np\ne"
        df.iloc[6,1] = "Bife ao molho de bacon e cenoura \nSopa Califórnia  (frango, batata, cenoura e abobrinha)\nAlmeirão e tomate\nHambúrguer de feijão branco ao sugo"
    elif caso == "04 Cardápio abril 2019 RU Versão Definitiva com resposta Blu símbolos.docx":
        df.iloc[1,6] = "07\nChá de erva-doce\nPão hot dog com \nDoce de abóbora/Salada de frutas\n"
    elif caso == "Cardápio abril 2018 FINAL USAR ESTE.docx":
        df.drop("DOMINGO.1",axis=1,inplace=True)
        df.iloc[3,1] = "Frango ao m. sugo– assado com alecrim (RUC)\nSopa de feijão\nEscarolae cenoura ralada\nPanqueca de grão-de-bico"
        df.iloc[11,0] = "23\nChá de hortelã\nPão hotdog com doce\nBanana/Melão\n"
        df.iloc[12,0] = "Guisadinho (moída c/ milho e ervilha)\nEspaguete ao alho e óleo Maçã\nNaN\nRepolho bicolor e cenoura ral.\nGuisadinho de PTS (moída c/ milho)"
        df.iloc[15,6] = "Barreado\nSopa abóbora\nNaN\nAcelga e tomate\nLentilha c/ legumes"
    elif caso == "01 Cardápio Março 2018 blum. revisado final.docx":
        df.iloc[3,4] = "Quibe \nSopa juliana\nAcelga e beterraba ralada\nQuibe de PTS"
        df.iloc[5,5] = "Filé de frango ao molho de mostarda\nLegumes refogados\nPudim de chocolate\nAlmeirão e tomate\nBOLINHO DE LENTILHA"
        df.iloc[9,4] = "Peixe à milanesa c/ limão\nPirão de peixe\nBanana\nAlface americana e cenoura ral.\nHambúrguer de soja"
        df.iloc[10,2] = "Bife à rolê\nSopa de mandioca\nAcelga e beterraba ralada\nBerinjela rolê"
        df.iloc[11,4] = "30\nChá de hortelã \nPão francês com doce de abóbora \nMelão"
        df.iloc[12,4] = "PEIXE NO FUBÁ AO MOLHO TÁRTARO\nBATATA AMASSADA\nBanana \nAcelga e beterraba ralada\nCreme de ervilha"
        df.iloc[13,5] = "Empadão de frango\nCaldo italiano\nAlface e rabanete\nEmpadão de legumes c/ feijão branco"
    elif caso == "Cardápio de setembro 2018 revisão almoço e jantar2.docx":
        df.iloc[8,3] = "20\nChá de erva-cidreira\nPão hot dog com doce de banana\nMaca"
    elif caso == "cardápio junho 2018 RU Revisado almoço Lu e vaneise.docx":
        df.iloc[2,4] = "Frango ao molho sugo (coxa/sob)\nPolenta cremosa\nMamão picado\nAlface crespa e pepino\nFeijão branco ao sugo"
    elif caso == "Cardápio de julho 2018 almoço e jantar revisado.docx":
        df.drop("DOMINGO.1",axis=1,inplace=True)
        df.iloc[1,0] = "02\nChá de hortelã\nPão hotdog com queijo\nMaçã/Ponkan"
        df.iloc[6,3] = "Panqueca de Carne moída\nCaldo Italiano\nRúcula e mista (cenoura, vagem, ab ) cozida \nPanqueca de P.T.S"
        df.iloc[12,4] = "Coxa/sobrecoxa ao molho curry\nSeleta de legumes\nNaN\nAlface crespa e pepino\nHambúrguer de soja"
        df.iloc[15,0] = "Picadinho de carne\nAbobrinha colorida \nSagú de vinho\nAcelga e cenoura ralada\nFeijão cavalo refogado"
        df.iloc[15,6] = "Strogonoff de carne\nPurê de batata\nNaN\nAcelga e tomate\nGrão de bico com champignon"
    elif caso == "Cardápio novembro 18 - revisado almoço e jantar FINAL 3.docx":
        df.drop("DOMINGO.1",axis=1,inplace=True)
        df.iloc[12,0] = "Guisadinho (moída c/ milho e ervilha)\nEspaguete ao alho e óleo\nMaçã\nRepolho bicolor e cenoura ralada\nGuisadinho de PTS c/ milho e ervilha"
        df.iloc[14,1] = "27\nChá de camomila\nPão francês com presunto\nMelancia"
        df.iloc[15,6] = "Strogonoff\nBatata chips\nNaN\nAcelga e tomate\nCharutinho de repolho com lentilha"
    elif caso == "Cardápio outubro 2018 FINAL definitivo.docx":
        df.iloc[1,6] = "07\nChá mate\nPão hot dog com doce de abóbora\nMelão"
        df.iloc[5,4] = "12\nChá de hortelã\nPãofrancês com presunto\nAbacaxi"
        df.iloc[7,3] = "Panqueca de carne\nCaldo italiano\nRadite e Chuchu c/ milho\nPanqueca(cenoura e ervilha seca) (era de brócolis, mas já tem no dia anterior)\n"
    elif caso == "cardápio maio RU revisado Lucyanne vaneise.docx":
        df.iloc[5,2] = "Carne moída refogada\nBolinho de cenoura\nPudim Prestígio\nEscarola e abobrinha\nTomate recheado c/ feijão branco temperado com cebola e tempero verde\n"
        df.iloc[5,3] = "Peixe frito no fubá c/ limão\nJardineira de legumes\nGelatina c/ creme\nAcelga e beterraba ralada\nBolinho de lentilha\n"
        df.iloc[5,4] = "Strogonoff de carne\nBatata chips\nLaranja\nAlface americana e tomate\nPimentão recheado com PTS, arroz e tomate\n"
        df.iloc[6,3] = "Bife à role\nSopa de batata-salsa\nAlface americana e vagem\nBerinjela à rolêcom pasta de feijão branco\n"
        df.iloc[9,1] = "Almôndega à chinesa\nCreme de espinafre/milho\nGelatina de cereja\nAgrião e rabanete\nQuibe de PTS\n"
        df.iloc[9,4] = "Picadinho de carne\nBolinho de arroz\nSagu de vinho\nAcelga e abobrinha ralada\nPicadinho vegano (pts ao sugo temperada)\n"
    return df
